import io
import os
import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document

# Templates are stored in the Next.js components directory
# Path relative from this file: backend/app/services/ → up 4 levels → src/app/dashboard/surat/components/
TEMPLATES_DIR = (
    Path(__file__).resolve().parent.parent.parent.parent
    / "src" / "app" / "dashboard" / "surat" / "components"
)

# LibreOffice executable candidates
LIBREOFFICE_PATHS = [
    "libreoffice",
    "soffice",
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/bin/libreoffice",
    "/usr/bin/soffice",
]


def _replace_in_paragraph(para, values: dict) -> None:
    """Replace {{key}} placeholders in a paragraph, handling split runs."""
    full_text = "".join(run.text for run in para.runs)
    if "{{" not in full_text:
        return

    def replacer(m: re.Match) -> str:
        key = m.group(1).strip()
        return str(values.get(key, ""))

    new_text = re.sub(r"\{\{(\w+)\}\}", replacer, full_text)
    if new_text == full_text:
        return

    # Write new text into first run, clear the rest to avoid duplication
    if para.runs:
        para.runs[0].text = new_text
        for run in para.runs[1:]:
            run.text = ""


def fill_template(template_filename: str, values: dict) -> bytes:
    """Load a docx template, replace {{field}} placeholders, return filled docx bytes."""
    template_path = TEMPLATES_DIR / template_filename

    if not template_path.exists():
        raise FileNotFoundError(
            f"Template '{template_filename}' tidak ditemukan di {TEMPLATES_DIR}"
        )

    doc = Document(str(template_path))

    # Paragraphs in body
    for para in doc.paragraphs:
        _replace_in_paragraph(para, values)

    # Tables in body
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, values)

    # Headers & footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            _replace_in_paragraph(para, values)
        for para in section.footer.paragraphs:
            _replace_in_paragraph(para, values)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def trim_pdf_pages(pdf_bytes: bytes, keep_pages: int) -> bytes:
    """Return a PDF containing only the first `keep_pages` pages."""
    import fitz  # pymupdf
    src = fitz.open(stream=pdf_bytes, filetype="pdf")
    if src.page_count <= keep_pages:
        src.close()
        return pdf_bytes
    dst = fitz.open()
    dst.insert_pdf(src, from_page=0, to_page=keep_pages - 1)
    result = dst.tobytes(garbage=4, deflate=True)
    src.close()
    dst.close()
    return result


def convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert docx bytes to PDF using LibreOffice headless. Raises RuntimeError if unavailable."""
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "surat.docx")
        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        converted = False
        for lo in LIBREOFFICE_PATHS:
            try:
                result = subprocess.run(
                    [lo, "--headless", "--convert-to", "pdf", "--outdir", tmpdir, docx_path],
                    capture_output=True,
                    timeout=30,
                )
                if result.returncode == 0:
                    converted = True
                    break
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue

        if not converted:
            raise RuntimeError(
                "LibreOffice tidak ditemukan. Install LibreOffice untuk konversi PDF, "
                "atau buka file .docx yang diunduh dengan Microsoft Word / LibreOffice."
            )

        pdf_path = os.path.join(tmpdir, "surat.pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError("Konversi PDF gagal — file output tidak terbuat.")

        with open(pdf_path, "rb") as f:
            return f.read()
